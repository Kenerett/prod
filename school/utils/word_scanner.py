# school/utils/word_scanner.py
"""
Утилиты для сканирования Word документов и извлечения предметов
Можно использовать как в командах управления, так и в представлениях
"""

import docx
import re
import os
from typing import List, Dict, Optional
import logging

logger = logging.getLogger(__name__)

class WordSubjectScanner:
    """Класс для сканирования Word документов и поиска названий предметов"""
    
    def __init__(self):
        self.patterns = {
            'course_unit_title': [
                r'course\s+unit\s+title',
                r'course\s+title',
                r'unit\s+title',
                r'предмет',
                r'дисциплина',
                r'название\s+предмета',
                r'название\s+дисциплины',
                r'course\s+name',
                r'subject\s+name'
            ],
            'extract_subject': [
                r'course\s+unit\s+title[:\s-]+(.+)',
                r'course\s+title[:\s-]+(.+)',
                r'unit\s+title[:\s-]+(.+)',
                r'предмет[:\s-]+(.+)',
                r'дисциплина[:\s-]+(.+)',
                r'название\s+предмета[:\s-]+(.+)',
                r'название\s+дисциплины[:\s-]+(.+)',
                r'course\s+name[:\s-]+(.+)',
                r'subject\s+name[:\s-]+(.+)'
            ]
        }
    
    def scan_document(self, file_path: str) -> List[Dict[str, str]]:
        """
        Сканирует Word документ и возвращает список найденных предметов
        
        Args:
            file_path (str): Путь к .docx файлу
            
        Returns:
            List[Dict]: Список словарей с информацией о предметах
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        if not file_path.lower().endswith('.docx'):
            raise ValueError("Поддерживаются только файлы .docx")
        
        try:
            doc = docx.Document(file_path)
            subjects = []
            
            # Сканируем параграфы
            subjects.extend(self._scan_paragraphs(doc.paragraphs))
            
            # Сканируем таблицы
            for table in doc.tables:
                subjects.extend(self._scan_table(table))
            
            # Убираем дубликаты
            unique_subjects = self._remove_duplicates(subjects)
            
            logger.info(f"Найдено {len(unique_subjects)} уникальных предметов в {file_path}")
            return unique_subjects
            
        except Exception as e:
            logger.error(f"Ошибка сканирования документа {file_path}: {str(e)}")
            raise
    
    def scan_directory(self, directory_path: str) -> List[Dict[str, str]]:
        """
        Сканирует все .docx файлы в директории
        
        Args:
            directory_path (str): Путь к директории
            
        Returns:
            List[Dict]: Список всех найденных предметов
        """
        if not os.path.exists(directory_path):
            raise FileNotFoundError(f"Директория не найдена: {directory_path}")
        
        docx_files = []
        for file in os.listdir(directory_path):
            if file.lower().endswith('.docx') and not file.startswith('~'):
                docx_files.append(os.path.join(directory_path, file))
        
        if not docx_files:
            logger.warning(f"В директории {directory_path} не найдены .docx файлы")
            return []
        
        all_subjects = []
        processed_files = 0
        
        for file_path in docx_files:
            try:
                subjects = self.scan_document(file_path)
                all_subjects.extend(subjects)
                processed_files += 1
                logger.info(f"Обработан файл: {os.path.basename(file_path)}")
                
            except Exception as e:
                logger.error(f"Ошибка обработки файла {file_path}: {str(e)}")
        
        unique_subjects = self._remove_duplicates(all_subjects)
        logger.info(f"Обработано {processed_files} файлов, найдено {len(unique_subjects)} уникальных предметов")
        
        return unique_subjects
    
    def _scan_paragraphs(self, paragraphs) -> List[Dict[str, str]]:
        """Сканирует параграфы документа"""
        subjects = []
        
        for i, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Проверяем, содержит ли параграф указатель на название предмета
            if self._contains_course_unit_title(text):
                # Пытаемся извлечь название из того же параграфа
                subject_info = self._extract_subject_from_text(text)
                if subject_info:
                    subjects.append(subject_info)
                    continue
                
                # Если в том же параграфе нет названия, ищем в следующих параграфах
                for j in range(1, min(4, len(paragraphs) - i)):  # Проверяем до 3 следующих параграфов
                    next_text = paragraphs[i + j].text.strip()
                    if next_text and not self._is_service_text(next_text):
                        subject_name = self._clean_subject_name(next_text)
                        if subject_name and len(subject_name) > 2:
                            subjects.append({
                                'name': subject_name,
                                'description': f'Предмет из документа: {subject_name}',
                                'source_text': text[:100],
                                'credits': self._extract_credits(next_text)
                            })
                            break
        
        return subjects
    
    def _scan_table(self, table) -> List[Dict[str, str]]:
        """Сканирует таблицу в поисках предметов"""
        subjects = []
        
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if not text:
                    continue
                
                if self._contains_course_unit_title(text):
                    # Ищем название предмета в той же ячейке
                    subject_info = self._extract_subject_from_text(text)
                    if subject_info:
                        subjects.append(subject_info)
                        continue
                    
                    # Ищем в соседних ячейках (справа, снизу)
                    found = False
                    
                    # Проверяем ячейки справа
                    for next_cell_idx in range(cell_idx + 1, len(row.cells)):
                        next_cell_text = row.cells[next_cell_idx].text.strip()
                        if next_cell_text and not self._is_service_text(next_cell_text):
                            subject_name = self._clean_subject_name(next_cell_text)
                            if subject_name and len(subject_name) > 2:
                                subjects.append({
                                    'name': subject_name,
                                    'description': f'Предмет из таблицы: {subject_name}',
                                    'source_text': text[:100],
                                    'credits': self._extract_credits(next_cell_text)
                                })
                                found = True
                                break
                    
                    # Если не найдено справа, проверяем снизу
                    if not found and row_idx + 1 < len(table.rows):
                        try:
                            below_cell = table.rows[row_idx + 1].cells[cell_idx]
                            below_text = below_cell.text.strip()
                            if below_text and not self._is_service_text(below_text):
                                subject_name = self._clean_subject_name(below_text)
                                if subject_name and len(subject_name) > 2:
                                    subjects.append({
                                        'name': subject_name,
                                        'description': f'Предмет из таблицы: {subject_name}',
                                        'source_text': text[:100],
                                        'credits': self._extract_credits(below_text)
                                    })
                        except IndexError:
                            pass  # Ячейка недоступна
        
        return subjects
    
    def _contains_course_unit_title(self, text: str) -> bool:
        """Проверяет, содержит ли текст указатель на название предмета"""
        text_lower = text.lower()
        for pattern in self.patterns['course_unit_title']:
            if re.search(pattern, text_lower):
                return True
        return False
    
    def _extract_subject_from_text(self, text: str) -> Optional[Dict[str, str]]:
        """Извлекает название предмета из текста"""
        for pattern in self.patterns['extract_subject']:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                raw_name = match.group(1).strip()
                subject_name = self._clean_subject_name(raw_name)
                
                if subject_name and len(subject_name) > 2 and not self._is_service_text(subject_name):
                    return {
                        'name': subject_name,
                        'description': f'Предмет: {subject_name}',
                        'source_text': text[:100],
                        'credits': self._extract_credits(raw_name)
                    }
        return None
    
    def _clean_subject_name(self, name: str) -> str:
        """Очищает название предмета от лишних символов"""
        if not name:
            return ""
        
        # Убираем лишние пробелы
        name = re.sub(r'\s+', ' ', name.strip())
        
        # Убираем лишние символы в начале и конце
        name = re.sub(r'^[:\-\s\._]+|[:\-\s\._]+$', '', name)
        
        # Убираем общие лишние фразы
        cleanup_patterns = [
            r'^(название|name|title|предмет|дисциплина|course|subject)[:\s]*',
            r'\([^)]*\)$',  # Убираем скобки в конце (кроме кредитов)
            r'^\d+\.\s*',   # Убираем номера в начале
            r'^-+\s*',      # Убираем тире в начале
        ]
        
        for pattern in cleanup_patterns:
            name = re.sub(pattern, '', name, flags=re.IGNORECASE).strip()
        
        # Ограничиваем длину
        if len(name) > 200:
            name = name[:200].strip()
        
        return name
    
    def _extract_credits(self, text: str) -> int:
        """Извлекает количество кредитов из текста"""
        # Ищем паттерны типа "3 credits", "(5 кредитов)", "ECTS: 4"
        patterns = [
            r'(\d+)\s*credits',
            r'(\d+)\s*кредитов?',
            r'(\d+)\s*ects',
            r'credits?\s*:\s*(\d+)',
            r'кредитов?\s*:\s*(\d+)',
            r'ects\s*:\s*(\d+)',
            r'\((\d+)\s*кредитов?\)',
            r'\((\d+)\s*credits?\)',
        ]
        
        text_lower = text.lower()
        for pattern in patterns:
            match = re.search(pattern, text_lower)
            if match:
                try:
                    return int(match.group(1))
                except (ValueError, IndexError):
                    continue
        
        return 0  # По умолчанию
    
    def _is_service_text(self, text: str) -> bool:
        """Проверяет, является ли текст служебным"""
        if not text or len(text.strip()) < 3:
            return True
        
        text_lower = text.lower().strip()
        
        service_patterns = [
            r'^\d+$',  # Только цифры
            r'^[^\w\s]*$',  # Только символы пунктуации
            r'^(table|таблица|row|строка|column|колонка|page|страница)',
            r'^(yes|no|да|нет|true|false)$',
            r'^\s*$',  # Пустая строка
            r'^(na|n/a|не указано|not specified)$',
            r'^(total|итого|sum|сумма)$',
        ]
        
        for pattern in service_patterns:
            if re.match(pattern, text_lower):
                return True
        
        return False
    
    def _remove_duplicates(self, subjects: List[Dict[str, str]]) -> List[Dict[str, str]]:
        """Убирает дубликаты предметов"""
        seen_names = set()
        unique_subjects = []
        
        for subject in subjects:
            name = subject['name'].strip().lower()
            if name not in seen_names and name:
                seen_names.add(name)
                unique_subjects.append(subject)
        
        return unique_subjects


# Функция для быстрого использования
def scan_word_document(file_path: str) -> List[Dict[str, str]]:
    """
    Быстрая функция для сканирования одного Word документа
    
    Args:
        file_path (str): Путь к .docx файлу
        
    Returns:
        List[Dict]: Список найденных предметов
    """
    scanner = WordSubjectScanner()
    return scanner.scan_document(file_path)


def scan_word_directory(directory_path: str) -> List[Dict[str, str]]:
    """
    Быстрая функция для сканирования директории с Word документами
    
    Args:
        directory_path (str): Путь к директории
        
    Returns:
        List[Dict]: Список всех найденных предметов
    """
    scanner = WordSubjectScanner()
    return scanner.scan_directory(directory_path)